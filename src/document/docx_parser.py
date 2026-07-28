"""DOCX 解析器（基于 python-docx）"""
from __future__ import annotations
import uuid
from pathlib import Path

from ..base import BaseParser, Document
from ..registry import register


@register("parser", "docx")
class DocxParser(BaseParser):
    def parse(self, file_path: str) -> Document:
        import docx

        doc = docx.Document(file_path)
        content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return Document(
            id=str(uuid.uuid4()),
            content=content,
            source=str(file_path),
            metadata={"format": "docx", "filename": Path(file_path).name},
        )

    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]
